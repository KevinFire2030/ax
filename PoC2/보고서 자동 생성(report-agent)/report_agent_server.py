"""Report Auto-Generation Agent PoC (Web Demo)

- Input: title/keyword + URL (+ optional pasted body)
- Retrieval: find similar existing reports (13 docx now; scalable to 300)
- Generation: follow AX report rules + Samsung support-team tone
- Output: preview + DOCX download

Run:
  cd /d E:\ax\상생\PoC\보고서 자동 생성(report-agent)
  Copy-Item .\.env.template .\.env
  # set OPENAI_API_KEY in .env
  pip install -r requirements.txt
  python report_agent_server.py

Open:
  http://127.0.0.1:8010/
"""

import json
import os
import re
import uuid
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

import trafilatura
from dotenv import load_dotenv
from fastapi import FastAPI
from fastapi.responses import FileResponse, HTMLResponse
from pydantic import BaseModel

from openai import OpenAI

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


# --------------------
# Env / settings
# --------------------

load_dotenv()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "").strip()
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-5.2").strip()
OPENAI_TEMPERATURE = float(os.getenv("OPENAI_TEMPERATURE", "0.2"))
OPENAI_MAX_OUTPUT_TOKENS = int(os.getenv("OPENAI_MAX_OUTPUT_TOKENS", "2500"))
OPENAI_TIMEOUT_SECONDS = int(os.getenv("OPENAI_TIMEOUT_SECONDS", "60"))

REPORT_CORPUS_DIR = Path(os.getenv("REPORT_CORPUS_DIR", ".\\data\\existing_reports"))
OUTPUT_DIR = Path(os.getenv("OUTPUT_DIR", ".\\outputs"))
TOP_K = int(os.getenv("TOP_K", "5"))

USER_AGENT = os.getenv("USER_AGENT", "Mozilla/5.0 (AX PoC)")

BASE_DIR = Path(__file__).resolve().parent
WEB_DIR = BASE_DIR / "web"
PROMPT_PATH = BASE_DIR / "prompts" / "tone_and_report_json.md"

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


# --------------------
# Models
# --------------------

class GenerateRequest(BaseModel):
    title: str
    url: Optional[str] = ""
    pasted_body: Optional[str] = ""


# --------------------
# Corpus loading / retrieval
# --------------------

@dataclass
class ReportDoc:
    doc_id: str
    title: str
    path: str
    text: str


def read_docx_text(path: Path) -> str:
    d = docx.Document(str(path))
    paras = [p.text.strip() for p in d.paragraphs if p.text and p.text.strip()]
    return "\n".join(paras)


def load_corpus() -> List[ReportDoc]:
    docs: List[ReportDoc] = []
    if not REPORT_CORPUS_DIR.exists():
        return docs

    for p in sorted(REPORT_CORPUS_DIR.glob("*.docx")):
        try:
            text = read_docx_text(p)
        except Exception:
            continue

        # title heuristic: filename without extension OR first line
        first_line = text.splitlines()[0].strip() if text else ""
        title = p.stem
        if first_line and len(first_line) < 120:
            # often same as filename
            title = first_line

        docs.append(
            ReportDoc(
                doc_id=p.stem,
                title=title,
                path=str(p),
                text=text,
            )
        )

    return docs


CORPUS: List[ReportDoc] = load_corpus()
VECTORIZER = TfidfVectorizer(max_features=40000)
CORPUS_MATRIX = None


def ensure_index():
    global CORPUS_MATRIX
    if CORPUS_MATRIX is not None:
        return
    texts = [d.title + "\n" + d.text[:4000] for d in CORPUS]
    CORPUS_MATRIX = VECTORIZER.fit_transform(texts) if texts else None


def retrieve_similar(query: str, k: int = 5) -> List[Dict[str, Any]]:
    if not CORPUS:
        return []
    ensure_index()
    if CORPUS_MATRIX is None:
        return []
    qv = VECTORIZER.transform([query])
    sims = cosine_similarity(qv, CORPUS_MATRIX)[0]
    idxs = sims.argsort()[::-1][:k]
    out = []
    for i in idxs:
        out.append({
            "doc_id": CORPUS[i].doc_id,
            "title": CORPUS[i].title,
            "score": float(sims[i]),
            "path": CORPUS[i].path,
        })
    return out


def make_style_examples(retrieval: List[Dict[str, Any]], max_chars: int = 4000) -> str:
    chunks = []
    for r in retrieval[:3]:
        doc = next((d for d in CORPUS if d.doc_id == r["doc_id"]), None)
        if not doc:
            continue
        sample = doc.text[:1500]
        chunks.append(f"[예시: {doc.title}]\n{sample}")
    joined = "\n\n".join(chunks)
    return joined[:max_chars]


# --------------------
# Article extraction
# --------------------

def fetch_article_text(url: str) -> str:
    if not url:
        return ""
    downloaded = trafilatura.fetch_url(url)
    if not downloaded:
        return ""
    txt = trafilatura.extract(downloaded, include_comments=False, include_tables=False)
    return txt or ""


# --------------------
# LLM generation
# --------------------

def load_prompt_template() -> str:
    return PROMPT_PATH.read_text(encoding="utf-8")


def render_prompt(template: str, *, user_title: str, url: str, article_text: str, style_examples: str, extra_sources: str) -> str:
    prompt = template
    prompt = prompt.replace("{{USER_TITLE}}", user_title)
    prompt = prompt.replace("{{URL}}", url)
    prompt = prompt.replace("{{ARTICLE_TEXT}}", article_text)
    prompt = prompt.replace("{{STYLE_EXAMPLES}}", style_examples)
    prompt = prompt.replace("{{EXTRA_SOURCES}}", extra_sources)
    return prompt


def call_llm(prompt: str) -> Dict[str, Any]:
    if not OPENAI_API_KEY:
        raise RuntimeError("OPENAI_API_KEY is empty")

    client = OpenAI(api_key=OPENAI_API_KEY, timeout=OPENAI_TIMEOUT_SECONDS)
    resp = client.responses.create(
        model=OPENAI_MODEL,
        input=prompt,
        temperature=OPENAI_TEMPERATURE,
        max_output_tokens=OPENAI_MAX_OUTPUT_TOKENS,
    )
    return json.loads(resp.output_text)


# --------------------
# DOCX rendering
# --------------------

def add_title(doc: docx.Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(f" {text} ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run.bold = True
    run.underline = True
    run.font.size = Pt(20)
    run.font.name = "바탕체"


def add_section_header(doc: docx.Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(f"□ {text}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "바탕체"


def add_bullets(doc: docx.Document, bullets: List[str], indent_spaces: int = 3):
    indent = " " * indent_spaces
    for b in bullets or []:
        p = doc.add_paragraph()
        run = p.add_run(f"{indent}- {b}")
        run.font.size = Pt(14)
        run.font.name = "바탕체"


def build_docx(report: Dict[str, Any], out_path: Path):
    doc = docx.Document()

    add_title(doc, report.get("report_title", ""))

    # Section 1
    add_section_header(doc, f"소제목 ① : {report.get('headline_one_liner','')}")
    add_bullets(doc, report.get("section1_summary_bullets", []), indent_spaces=3)

    # Section 2
    add_section_header(doc, "소제목 ② : 추가 확인 사항")
    add_bullets(doc, report.get("section2_additional_checks_bullets", []), indent_spaces=3)

    # Insights
    add_section_header(doc, "주요 인사이트(시사점)")
    add_bullets(doc, report.get("insights_bullets", []), indent_spaces=3)

    # Appendix
    add_section_header(doc, "별첨")
    appendix = report.get("appendix", {}) or {}

    companies = appendix.get("companies", []) or []
    people = appendix.get("people", []) or []

    if companies:
        add_section_header(doc, "회사 개요")
        for c in companies:
            name = c.get("name", "")
            bullets = c.get("overview_bullets", []) or []
            add_bullets(doc, [f"{name}"] + bullets, indent_spaces=3)

    if people:
        add_section_header(doc, "인물 프로필")
        for p in people:
            name = p.get("name", "")
            bullets = p.get("profile_bullets", []) or []
            add_bullets(doc, [f"{name}"] + bullets, indent_spaces=3)

    doc.save(str(out_path))


# --------------------
# FastAPI
# --------------------

app = FastAPI(title="AX Report Auto-Generation PoC")


@app.get("/")
def index():
    return HTMLResponse((WEB_DIR / "index.html").read_text(encoding="utf-8"))


@app.post("/api/generate")
def api_generate(req: GenerateRequest):
    try:
        user_title = (req.title or "").strip()
        if not user_title:
            return {"ok": False, "error": "제목/키워드는 필수입니다."}

        url = (req.url or "").strip()
        pasted = (req.pasted_body or "").strip()

        retrieval = retrieve_similar(user_title, k=TOP_K)
        style_examples = make_style_examples(retrieval)

        article_text = ""
        if url:
            article_text = fetch_article_text(url)

        if not article_text:
            article_text = pasted

        if not article_text:
            return {"ok": False, "error": "기사 본문을 가져오지 못했습니다. 본문 붙여넣기를 사용하세요."}

        prompt_template = load_prompt_template()
        prompt = render_prompt(
            prompt_template,
            user_title=user_title,
            url=url,
            article_text=article_text[:12000],
            style_examples=style_examples,
            extra_sources="(PoC: 생략)",
        )

        report = call_llm(prompt)

        # write docx
        run_id = datetime.now().strftime("%Y%m%d_%H%M%S") + "_" + uuid.uuid4().hex[:8]
        out_docx = OUTPUT_DIR / f"AX_report_{run_id}.docx"
        build_docx(report, out_docx)

        return {
            "ok": True,
            "retrieval": retrieval,
            "report": report,
            "download_url": f"/api/download/{out_docx.name}",
        }

    except Exception as e:
        return {"ok": False, "error": str(e)}


@app.get("/api/download/{filename}")
def api_download(filename: str):
    p = OUTPUT_DIR / filename
    if not p.exists():
        return {"ok": False, "error": "file not found"}
    return FileResponse(str(p), filename=filename, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


if __name__ == "__main__":
    import uvicorn

    host = os.getenv("HOST", "127.0.0.1")
    port = int(os.getenv("PORT", "8010"))
    uvicorn.run(app, host=host, port=port)
